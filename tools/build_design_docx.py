from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "课程设计说明书_课题3_提交版.docx"

SOURCES = [
    ROOT / "docs" / "课程设计说明书_课题3_初稿.md",
    ROOT / "docs" / "需求实现矩阵.md",
    ROOT / "docs" / "汇编程序清单与模块说明.md",
    ROOT / "docs" / "仿真测试记录.md",
]


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
TABLE_FILL = "F2F4F7"
BODY_FONT = "宋体"
LATIN_FONT = "Calibri"
CODE_FONT = "Consolas"


def set_run_font(run, name=BODY_FONT, latin=LATIN_FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(p, before=0, after=6, line=1.25, first_line=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line is not None:
        fmt.first_line_indent = Cm(first_line)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=0, line=1.15)
    run = p.add_run(clean_inline(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_width(table, col_count):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    width = Inches(6.35)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width.inches * 1440)))

    col_width = width / max(col_count, 1)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width


def set_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=0, line=1.0)
    run = p.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)

    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    run = p.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    run = p.add_run("1")
    set_run_font(run, size=9, color=MUTED)

    run = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    run = p.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    set_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = LATIN_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.2

    code = styles.add_style("CodeBlock", 1)
    code.font.name = CODE_FONT
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    code.font.size = Pt(8.5)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(2)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.left_indent = Cm(0.45)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=42, after=8, line=1.2)
    run = p.add_run("《单片机原理及接口技术》课程设计说明书")
    set_run_font(run, size=22, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=28, line=1.2)
    run = p.add_run("课题 3：流水灯及智能小车超声测距及避障实现")
    set_run_font(run, size=16, color=DARK_BLUE, bold=True)

    table = doc.add_table(rows=6, cols=2)
    set_table_width(table, 2)
    rows = [
        ("学生学院", "请填写"),
        ("专业班级", "请填写"),
        ("小组成员", "请按贡献大小顺序填写姓名与学号"),
        ("指导教师", "请填写"),
        ("计划完成日期", "2026 年 6 月 26 日"),
        ("提交材料", "说明书、汇编程序、Proteus 仿真文件、20 秒内调试视频"),
    ]
    for i, (k, v) in enumerate(rows):
        set_cell_text(table.cell(i, 0), k, bold=True, size=10)
        set_cell_text(table.cell(i, 1), v, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=22, after=0, line=1.2)
    run = p.add_run("说明：封面署名顺序按个人贡献大小排列。")
    set_run_font(run, size=10, color=MUTED)
    doc.add_page_break()


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1（\2）", text)
    return text.strip()


def add_inline_runs(p, text: str):
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1（\2）", text)
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, name="等线", latin=CODE_FONT, size=10)
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=11, bold=True)
        else:
            run = p.add_run(part)
            set_run_font(run, size=11)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    set_table_width(table, col_count)
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            text = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(table.cell(r_idx, c_idx), text, bold=(r_idx == 0), size=8.5 if col_count >= 4 else 9)
            if r_idx == 0:
                shade_cell(table.cell(r_idx, c_idx), TABLE_FILL)
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=4, line=1.0)


def list_number_abstract_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    for abstract in numbering.findall(qn("w:abstractNum")):
        for p_style in abstract.iter(qn("w:pStyle")):
            if p_style.get(qn("w:val")) == "ListNumber":
                return int(abstract.get(qn("w:abstractNumId")))
    return 7


def new_numbered_list_id(doc: Document, start: int = 1) -> int:
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    num = numbering.add_num(list_number_abstract_id(doc))
    override = num.add_lvlOverride(ilvl=0)
    override.add_startOverride(start)
    return int(num.numId)


def apply_numbering(p, num_id: int):
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id


def add_markdown_file(doc: Document, path: Path, appendix=False):
    if appendix:
        doc.add_page_break()
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    i = 0
    in_code = False
    code_lines = []
    para_lines = []
    active_numbered_list_id = None

    def flush_para():
        nonlocal para_lines
        if not para_lines:
            return
        content = " ".join(x.strip() for x in para_lines if x.strip())
        para_lines = []
        if content:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, after=6, line=1.25, first_line=0.74)
            add_inline_runs(p, content)

    def reset_numbered_list():
        nonlocal active_numbered_list_id
        active_numbered_list_id = None

    def flush_code():
        nonlocal code_lines
        for code_line in code_lines:
            p = doc.add_paragraph(style="CodeBlock")
            p.add_run(code_line.rstrip())
        code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_para()
            reset_numbered_list()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            flush_para()
            reset_numbered_list()
            i += 1
            continue

        if stripped.startswith("|"):
            flush_para()
            reset_numbered_list()
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_para()
            reset_numbered_list()
            level = len(heading.group(1))
            title = clean_inline(heading.group(2))
            if level == 1:
                p = doc.add_heading(title, level=1)
            elif level == 2:
                p = doc.add_heading(title, level=2)
            else:
                p = doc.add_heading(title, level=3)
            for run in p.runs:
                set_run_font(run, size=run.font.size.pt if run.font.size else None, bold=True)
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^(\d+)[.)]\s+(.+)$", stripped)
        if bullet or numbered:
            flush_para()
            if bullet:
                reset_numbered_list()
            else:
                if active_numbered_list_id is None:
                    active_numbered_list_id = new_numbered_list_id(doc, int(numbered.group(1)))
            text = clean_inline(bullet.group(1) if bullet else numbered.group(2))
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            if numbered:
                apply_numbering(p, active_numbered_list_id)
            set_paragraph_spacing(p, after=4, line=1.2)
            add_inline_runs(p, text)
            i += 1
            continue

        reset_numbered_list()
        para_lines.append(line)
        i += 1

    flush_para()
    flush_code()


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)
    add_cover(doc)

    for idx, source in enumerate(SOURCES):
        add_markdown_file(doc, source, appendix=(idx > 0))

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
